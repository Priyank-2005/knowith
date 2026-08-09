"""
Knowith Capital — Document Generation Base Module
===================================================
Shared styling engine and reusable components for enterprise document generation.

This module provides the DocumentGenerator base class used by all four
document generators. It ensures consistent formatting, typography, colors,
and professional presentation across all deliverables.

Author: Strategy & Technology Consulting Division
Version: 1.0
"""

import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# ============================================================================
# CORPORATE COLOR PALETTE
# ============================================================================
class Colors:
    """Corporate color palette for Knowith Capital documentation."""

    # --- Primary Brand Colors ---
    NAVY = RGBColor(0x1B, 0x2A, 0x4A)
    NAVY_HEX = "1B2A4A"

    GOLD = RGBColor(0xC9, 0xA9, 0x6E)
    GOLD_HEX = "C9A96E"

    # --- Text Colors ---
    DARK_TEXT = RGBColor(0x2C, 0x3E, 0x50)
    BODY_TEXT = RGBColor(0x34, 0x3A, 0x40)
    LIGHT_TEXT = RGBColor(0x6C, 0x75, 0x7D)
    WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

    # --- Background Colors ---
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT_GRAY = RGBColor(0xF8, 0xF9, 0xFA)
    MEDIUM_GRAY = RGBColor(0xE9, 0xEC, 0xEF)
    DARK_GRAY = RGBColor(0x6C, 0x75, 0x7D)

    # --- Accent Colors ---
    ACCENT_BLUE = RGBColor(0x2E, 0x86, 0xAB)
    SUCCESS = RGBColor(0x28, 0xA7, 0x45)
    WARNING = RGBColor(0xFD, 0x7E, 0x14)
    DANGER = RGBColor(0xDC, 0x35, 0x45)
    INFO = RGBColor(0x17, 0xA2, 0xB8)
    PURPLE = RGBColor(0x7B, 0x1F, 0xA2)

    # --- Table Colors (hex strings for OxmlElement shading) ---
    TABLE_HEADER_BG = "1B2A4A"
    TABLE_ALT_ROW_BG = "F8F9FA"
    TABLE_BORDER_COLOR = "BFBFBF"

    # --- Callout Box Colors (hex strings) ---
    NOTE_BG = "EBF5FB"
    NOTE_BORDER = "2E86AB"
    NOTE_TITLE_COLOR = RGBColor(0x1A, 0x5C, 0x7A)

    REC_BG = "E8F5E9"
    REC_BORDER = "28A745"
    REC_TITLE_COLOR = RGBColor(0x1B, 0x7A, 0x30)

    WARN_BG = "FFF8E1"
    WARN_BORDER = "F57C00"
    WARN_TITLE_COLOR = RGBColor(0xE6, 0x51, 0x00)

    FUTURE_BG = "F3E5F5"
    FUTURE_BORDER = "7B1FA2"
    FUTURE_TITLE_COLOR = RGBColor(0x6A, 0x1B, 0x9A)

    IMPORTANT_BG = "FFF3E0"
    IMPORTANT_BORDER = "E65100"
    IMPORTANT_TITLE_COLOR = RGBColor(0xBF, 0x36, 0x0C)

    # --- Cover page colors ---
    COVER_LINE = "C9A96E"
    COVER_BG = "1B2A4A"


# ============================================================================
# TYPOGRAPHY CONSTANTS
# ============================================================================
class Fonts:
    """Font family definitions for consistent typography."""
    HEADING = "Calibri Light"
    BODY = "Calibri"
    COVER_TITLE = "Calibri Light"
    MONOSPACE = "Consolas"


class FontSizes:
    """Standardized font size scale across all documents."""
    # Cover page
    COVER_TITLE = Pt(32)
    COVER_SUBTITLE = Pt(15)
    COVER_META = Pt(11)
    COVER_CONFIDENTIAL = Pt(8.5)

    # Heading hierarchy
    H1 = Pt(22)
    H2 = Pt(16)
    H3 = Pt(13)
    H4 = Pt(11.5)

    # Body text
    BODY = Pt(10.5)
    SMALL = Pt(9)
    CAPTION = Pt(8.5)

    # Table text
    TABLE_HEADER = Pt(9.5)
    TABLE_BODY = Pt(9)

    # Header / Footer
    HEADER = Pt(8)
    FOOTER = Pt(8)

    # Callout boxes
    CALLOUT_TITLE = Pt(10.5)
    CALLOUT_BODY = Pt(10)


# ============================================================================
# SPACING CONSTANTS
# ============================================================================
class Spacing:
    """Standardized spacing values (in Points) for vertical rhythm."""
    BEFORE_H1 = Pt(28)
    AFTER_H1 = Pt(12)
    BEFORE_H2 = Pt(22)
    AFTER_H2 = Pt(8)
    BEFORE_H3 = Pt(16)
    AFTER_H3 = Pt(6)
    BEFORE_H4 = Pt(12)
    AFTER_H4 = Pt(4)

    BEFORE_BODY = Pt(2)
    AFTER_BODY = Pt(6)

    BEFORE_LIST = Pt(2)
    AFTER_LIST = Pt(3)

    TABLE_CELL_TOP = Pt(4)
    TABLE_CELL_BOTTOM = Pt(4)


# ============================================================================
# PROJECT INFORMATION CONSTANTS
# ============================================================================
class ProjectInfo:
    """Project metadata constants used across all documents."""

    # --- Client Details ---
    CLIENT_NAME = "Knowith Capital"
    CLIENT_INDUSTRY = "Investment / Mutual Fund Distribution"
    CLIENT_COUNTRY = "India"
    CLIENT_ENTITY_TYPE = "Mutual Fund Distributor (MFD)"

    # --- Project Details ---
    PROJECT_NAME = "Premium Website Design & Development"
    PROJECT_PHASE = "Phase 01 \u2014 Website"

    # --- Document Metadata ---
    PREPARED_BY = "Strategy & Technology Consulting Division"
    REVIEWED_BY = "Senior Leadership Team"
    APPROVED_BY = "Project Steering Committee"
    CONSULTING_FIRM = "Enterprise Solutions Group"

    VERSION = "1.0"
    STATUS = "Final Draft"
    CLASSIFICATION = "CONFIDENTIAL"

    # --- Notices ---
    CONFIDENTIALITY_NOTICE = (
        "This document contains confidential and proprietary information belonging to "
        "Knowith Capital. It is intended solely for the use of the individual or entity "
        "to whom it is addressed. Any unauthorized review, use, disclosure, or distribution "
        "is prohibited. If you are not the intended recipient, please contact the sender "
        "immediately and destroy all copies of the original document. This document must "
        "not be reproduced, in whole or in part, without the express written consent of "
        "Knowith Capital."
    )

    REGULATORY_NOTICE = (
        "Knowith Capital operates as a Mutual Fund Distributor (MFD) registered with "
        "the Association of Mutual Funds in India (AMFI). Knowith Capital is NOT a "
        "SEBI Registered Investment Advisor (RIA). All references to services within this "
        "document pertain exclusively to mutual fund distribution and related investor "
        "support services. No content in this document should be construed as investment "
        "advice, financial planning, or portfolio management services."
    )

    COPYRIGHT_NOTICE = (
        "\u00a9 {year} Knowith Capital. All rights reserved. No part of this document "
        "may be reproduced, distributed, or transmitted in any form or by any means "
        "without the prior written permission of Knowith Capital."
    )

    # --- Regulation-Safe Terminology ---
    SAFE_TERMS = {
        "service_type": "Mutual Fund Distribution",
        "service_action": "Investment Execution",
        "service_offering": "Wealth Solutions",
        "support": "Investment Support",
        "research": "Research & Insights",
        "communication": "Investor Communication",
        "role": "Mutual Fund Distributor",
        "entity": "Distribution Partner",
    }

    PROHIBITED_TERMS = [
        "Investment Advisor",
        "Financial Advisor",
        "Advisory Services",
        "Consultancy",
        "Investment Advice",
        "Financial Planning Services",
        "Portfolio Management Services",
        "Wealth Management Advisor",
    ]

    # --- Future Modules (Phase 02+) ---
    FUTURE_MODULES = [
        "Investor Portal / Dashboard",
        "AI Chat Assistant",
        "Portfolio Analyzer",
        "Financial Health Analyzer",
        "Research Assistant (AI)",
        "AI Email Assistant",
        "WhatsApp Assistant",
        "CRM System",
        "Admin Dashboard",
        "Analytics Dashboard",
    ]

    @classmethod
    def get_date(cls):
        """Return the current date formatted for document headers."""
        return datetime.datetime.now().strftime("%B %d, %Y")

    @classmethod
    def get_year(cls):
        """Return the current year."""
        return datetime.datetime.now().strftime("%Y")


# ============================================================================
# DOCUMENT GENERATOR BASE CLASS
# ============================================================================
class DocumentGenerator:
    """
    Base class for all Knowith Capital document generators.

    Provides shared functionality for creating professionally formatted
    Microsoft Word documents with consistent styling, typography, and layout.

    Usage:
        class MyDocGenerator(DocumentGenerator):
            def generate(self):
                self.add_cover_page("My Document", "Subtitle")
                self.add_front_matter()
                self.add_heading_1("Section Title")
                self.add_body_text("Content goes here.")
                self.save("output.docx")
    """

    def __init__(self, output_dir="documents"):
        """
        Initialize the document generator.

        Args:
            output_dir: Directory path for saving generated documents.
        """
        self.doc = Document()
        self.output_dir = output_dir
        self._configure_default_styles()
        self._configure_page_layout()

    # ========================================================================
    # PRIVATE: STYLE CONFIGURATION
    # ========================================================================

    def _configure_default_styles(self):
        """Configure all default document styles for consistent formatting."""
        style = self.doc.styles

        # --- Normal (body) style ---
        normal = style['Normal']
        normal.font.name = Fonts.BODY
        normal.font.size = FontSizes.BODY
        normal.font.color.rgb = Colors.BODY_TEXT
        normal.paragraph_format.space_after = Spacing.AFTER_BODY
        normal.paragraph_format.space_before = Spacing.BEFORE_BODY
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.3

        # --- Heading 1 ---
        h1 = style['Heading 1']
        h1.font.name = Fonts.HEADING
        h1.font.size = FontSizes.H1
        h1.font.color.rgb = Colors.NAVY
        h1.font.bold = True
        h1.paragraph_format.space_before = Spacing.BEFORE_H1
        h1.paragraph_format.space_after = Spacing.AFTER_H1
        h1.paragraph_format.keep_with_next = True

        # --- Heading 2 ---
        h2 = style['Heading 2']
        h2.font.name = Fonts.HEADING
        h2.font.size = FontSizes.H2
        h2.font.color.rgb = Colors.NAVY
        h2.font.bold = True
        h2.paragraph_format.space_before = Spacing.BEFORE_H2
        h2.paragraph_format.space_after = Spacing.AFTER_H2
        h2.paragraph_format.keep_with_next = True

        # --- Heading 3 ---
        h3 = style['Heading 3']
        h3.font.name = Fonts.HEADING
        h3.font.size = FontSizes.H3
        h3.font.color.rgb = Colors.DARK_TEXT
        h3.font.bold = True
        h3.paragraph_format.space_before = Spacing.BEFORE_H3
        h3.paragraph_format.space_after = Spacing.AFTER_H3
        h3.paragraph_format.keep_with_next = True

        # --- Heading 4 ---
        h4 = style['Heading 4']
        h4.font.name = Fonts.HEADING
        h4.font.size = FontSizes.H4
        h4.font.color.rgb = Colors.DARK_TEXT
        h4.font.bold = True
        h4.font.italic = True
        h4.paragraph_format.space_before = Spacing.BEFORE_H4
        h4.paragraph_format.space_after = Spacing.AFTER_H4
        h4.paragraph_format.keep_with_next = True

        # --- List Bullet ---
        if 'List Bullet' in style:
            lb = style['List Bullet']
            lb.font.name = Fonts.BODY
            lb.font.size = FontSizes.BODY
            lb.font.color.rgb = Colors.BODY_TEXT
            lb.paragraph_format.space_before = Spacing.BEFORE_LIST
            lb.paragraph_format.space_after = Spacing.AFTER_LIST

    def _configure_page_layout(self):
        """Configure page size, margins, and orientation."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)   # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # ========================================================================
    # COVER PAGE
    # ========================================================================

    def add_cover_page(self, doc_title, doc_subtitle=""):
        """
        Create a professional cover page with corporate styling.

        Args:
            doc_title:    Primary document title.
            doc_subtitle: Secondary subtitle or description.
        """
        # --- Top spacer ---
        for _ in range(4):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.space_before = Pt(0)

        # --- Gold accent line ---
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_para.paragraph_format.space_after = Pt(16)
        run = line_para.add_run("\u2500" * 40)
        run.font.color.rgb = Colors.GOLD
        run.font.size = Pt(14)

        # --- Document Classification ---
        cls_para = self.doc.add_paragraph()
        cls_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cls_para.paragraph_format.space_after = Pt(20)
        run = cls_para.add_run(ProjectInfo.CLASSIFICATION)
        run.font.name = Fonts.BODY
        run.font.size = Pt(10)
        run.font.color.rgb = Colors.GOLD
        run.font.bold = True
        run.font.all_caps = True

        # --- Document Title ---
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(8)
        run = title_para.add_run(doc_title)
        run.font.name = Fonts.COVER_TITLE
        run.font.size = FontSizes.COVER_TITLE
        run.font.color.rgb = Colors.NAVY
        run.font.bold = True

        # --- Document Subtitle ---
        if doc_subtitle:
            sub_para = self.doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_para.paragraph_format.space_after = Pt(24)
            run = sub_para.add_run(doc_subtitle)
            run.font.name = Fonts.HEADING
            run.font.size = FontSizes.COVER_SUBTITLE
            run.font.color.rgb = Colors.DARK_GRAY

        # --- Gold accent line ---
        line_para2 = self.doc.add_paragraph()
        line_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_para2.paragraph_format.space_after = Pt(30)
        run = line_para2.add_run("\u2500" * 40)
        run.font.color.rgb = Colors.GOLD
        run.font.size = Pt(14)

        # --- Metadata Block ---
        meta_items = [
            ("Client", ProjectInfo.CLIENT_NAME),
            ("Project", ProjectInfo.PROJECT_NAME),
            ("Phase", ProjectInfo.PROJECT_PHASE),
            ("Prepared By", ProjectInfo.PREPARED_BY),
            ("Version", ProjectInfo.VERSION),
            ("Status", ProjectInfo.STATUS),
            ("Date", ProjectInfo.get_date()),
        ]
        for label, value in meta_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            # Label
            lbl_run = p.add_run(f"{label}:  ")
            lbl_run.font.name = Fonts.BODY
            lbl_run.font.size = FontSizes.COVER_META
            lbl_run.font.color.rgb = Colors.LIGHT_TEXT
            lbl_run.font.bold = False
            # Value
            val_run = p.add_run(value)
            val_run.font.name = Fonts.BODY
            val_run.font.size = FontSizes.COVER_META
            val_run.font.color.rgb = Colors.DARK_TEXT
            val_run.font.bold = True

        # --- Spacer before confidentiality ---
        self.doc.add_paragraph().paragraph_format.space_after = Pt(20)

        # --- Confidentiality Notice ---
        conf_para = self.doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        conf_para.paragraph_format.space_before = Pt(12)
        run = conf_para.add_run("Confidentiality Notice")
        run.font.name = Fonts.BODY
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.NAVY
        run.font.bold = True

        notice_para = self.doc.add_paragraph()
        notice_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = notice_para.add_run(ProjectInfo.CONFIDENTIALITY_NOTICE)
        run.font.name = Fonts.BODY
        run.font.size = FontSizes.COVER_CONFIDENTIAL
        run.font.color.rgb = Colors.LIGHT_TEXT
        run.font.italic = True

        # --- Page break after cover ---
        self.doc.add_page_break()

    # ========================================================================
    # FRONT MATTER (Version History, Document Control, TOC)
    # ========================================================================

    def add_front_matter(self, versions=None):
        """
        Add standard front matter: version history, document control,
        regulatory notice, and table of contents.

        Args:
            versions: List of tuples (version, date, author, description).
                      Defaults to a single initial version entry.
        """
        self._add_version_history(versions)
        self.doc.add_page_break()
        self._add_document_control()
        self.doc.add_page_break()
        self._add_regulatory_notice()
        self.doc.add_page_break()
        self._add_table_of_contents()
        self.doc.add_page_break()

    def _add_version_history(self, versions=None):
        """Add a version history table."""
        self.add_heading_1("Version History")

        if versions is None:
            versions = [
                ("1.0", ProjectInfo.get_date(),
                 ProjectInfo.PREPARED_BY,
                 "Initial document creation \u2014 Final Draft"),
            ]

        headers = ["Version", "Date", "Author", "Description"]
        rows = [list(v) for v in versions]
        self.add_professional_table(headers, rows)

    def _add_document_control(self):
        """Add document control section with approval matrix."""
        self.add_heading_1("Document Control")

        # --- Document Information ---
        self.add_heading_2("Document Information")
        info_data = [
            ("Document Title", ""),
            ("Document Reference", ""),
            ("Version", ProjectInfo.VERSION),
            ("Status", ProjectInfo.STATUS),
            ("Classification", ProjectInfo.CLASSIFICATION),
            ("Date", ProjectInfo.get_date()),
            ("Client", ProjectInfo.CLIENT_NAME),
            ("Prepared By", ProjectInfo.PREPARED_BY),
        ]
        self.add_info_table(info_data)

        # --- Approval Matrix ---
        self.add_heading_2("Approval Matrix")
        approval_headers = ["Role", "Name", "Signature", "Date"]
        approval_rows = [
            ["Document Author", "", "", ""],
            ["Technical Reviewer", "", "", ""],
            ["Project Manager", "", "", ""],
            ["Client Stakeholder", "", "", ""],
        ]
        self.add_professional_table(approval_headers, approval_rows)

        # --- Distribution List ---
        self.add_heading_2("Distribution List")
        dist_headers = ["Name", "Role", "Organization", "Copy Type"]
        dist_rows = [
            ["", "Project Sponsor", ProjectInfo.CLIENT_NAME, "Electronic"],
            ["", "Project Manager", ProjectInfo.CLIENT_NAME, "Electronic"],
            ["", "Technical Lead", "Development Partner", "Electronic"],
            ["", "Design Lead", "Design Partner", "Electronic"],
            ["", "Content Strategist", "Content Partner", "Electronic"],
        ]
        self.add_professional_table(dist_headers, dist_rows)

    def _add_regulatory_notice(self):
        """Add regulatory compliance notice."""
        self.add_heading_1("Regulatory Notice")
        self.add_callout_box(
            "Important Regulatory Disclaimer",
            ProjectInfo.REGULATORY_NOTICE,
            box_type="important"
        )

        self.add_body_text(
            "Throughout this document, all references to services provided by "
            "Knowith Capital are limited to mutual fund distribution activities. "
            "The following terminology guidelines are observed:"
        )

        # Safe terms
        self.add_heading_3("Approved Terminology")
        for key, term in ProjectInfo.SAFE_TERMS.items():
            label = key.replace("_", " ").title()
            self.add_bullet_point(f"{label}: {term}")

        # Prohibited terms
        self.add_heading_3("Prohibited Terminology")
        self.add_body_text(
            "The following terms must NOT be used in any client-facing content, "
            "marketing materials, or website copy:"
        )
        for term in ProjectInfo.PROHIBITED_TERMS:
            self.add_bullet_point(term)

    def _add_table_of_contents(self):
        """Insert a Word TOC field code that auto-updates on document open."""
        self.add_heading_1("Table of Contents")

        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)

        # Begin field
        run1 = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run1._r.append(fld_begin)

        # Field instruction
        run2 = paragraph.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instr)

        # Separate
        run3 = paragraph.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fld_sep)

        # Placeholder text
        run4 = paragraph.add_run(
            '[Right-click here and select "Update Field" to generate the Table of Contents]'
        )
        run4.font.color.rgb = Colors.LIGHT_TEXT
        run4.font.italic = True
        run4.font.size = FontSizes.BODY

        # End field
        run5 = paragraph.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fld_end)

    # ========================================================================
    # HEADERS AND FOOTERS
    # ========================================================================

    def add_headers_footers(self, doc_title):
        """
        Add professional headers and footers to all sections.

        Args:
            doc_title: Title text displayed in the header.
        """
        for section in self.doc.sections:
            # --- Differentiate first page ---
            section.different_first_page_header_footer = True

            # --- Header (non-first pages) ---
            header = section.header
            header.is_linked_to_previous = False
            h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_para.paragraph_format.space_after = Pt(4)

            # Client name
            run_client = h_para.add_run(f"{ProjectInfo.CLIENT_NAME}  |  ")
            run_client.font.name = Fonts.BODY
            run_client.font.size = FontSizes.HEADER
            run_client.font.color.rgb = Colors.NAVY
            run_client.font.bold = True

            # Doc title
            run_title = h_para.add_run(doc_title)
            run_title.font.name = Fonts.BODY
            run_title.font.size = FontSizes.HEADER
            run_title.font.color.rgb = Colors.LIGHT_TEXT
            run_title.font.bold = False

            # Header bottom border
            self._add_paragraph_border_bottom(h_para, Colors.GOLD_HEX)

            # --- Footer (non-first pages) ---
            footer = section.footer
            footer.is_linked_to_previous = False
            f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            f_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            f_para.paragraph_format.space_before = Pt(6)

            # Footer top border
            self._add_paragraph_border_top(f_para, Colors.GOLD_HEX)

            # Confidential label
            run_conf = f_para.add_run(f"{ProjectInfo.CLASSIFICATION}  |  ")
            run_conf.font.name = Fonts.BODY
            run_conf.font.size = FontSizes.FOOTER
            run_conf.font.color.rgb = Colors.LIGHT_TEXT
            run_conf.font.bold = False

            # Copyright
            year = ProjectInfo.get_year()
            run_copy = f_para.add_run(f"\u00a9 {year} {ProjectInfo.CLIENT_NAME}")
            run_copy.font.name = Fonts.BODY
            run_copy.font.size = FontSizes.FOOTER
            run_copy.font.color.rgb = Colors.LIGHT_TEXT

            # Page number (right-aligned via tab)
            run_tab = f_para.add_run("\t\t\t\t\t")
            run_page_label = f_para.add_run("Page ")
            run_page_label.font.name = Fonts.BODY
            run_page_label.font.size = FontSizes.FOOTER
            run_page_label.font.color.rgb = Colors.LIGHT_TEXT

            # Page number field
            self._add_page_number_field(f_para)

    def _add_page_number_field(self, paragraph):
        """Insert a Word PAGE field code for automatic page numbering."""
        run = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        run2 = paragraph.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run2._r.append(instr)
        run2.font.name = Fonts.BODY
        run2.font.size = FontSizes.FOOTER
        run2.font.color.rgb = Colors.LIGHT_TEXT

        run3 = paragraph.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fld_sep)

        run4 = paragraph.add_run("1")
        run4.font.name = Fonts.BODY
        run4.font.size = FontSizes.FOOTER
        run4.font.color.rgb = Colors.LIGHT_TEXT

        run5 = paragraph.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fld_end)

    def _add_paragraph_border_bottom(self, paragraph, color_hex):
        """Add a bottom border to a paragraph element."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_paragraph_border_top(self, paragraph, color_hex):
        """Add a top border to a paragraph element."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '4')
        top.set(qn('w:color'), color_hex)
        pBdr.append(top)
        pPr.append(pBdr)

    # ========================================================================
    # HEADING METHODS
    # ========================================================================

    def add_heading_1(self, text):
        """Add a Heading 1 paragraph. Used for major document sections."""
        heading = self.doc.add_heading(text, level=1)
        return heading

    def add_heading_2(self, text):
        """Add a Heading 2 paragraph. Used for subsections."""
        heading = self.doc.add_heading(text, level=2)
        return heading

    def add_heading_3(self, text):
        """Add a Heading 3 paragraph. Used for topic-level items."""
        heading = self.doc.add_heading(text, level=3)
        return heading

    def add_heading_4(self, text):
        """Add a Heading 4 paragraph. Used for fine-grained sub-topics."""
        heading = self.doc.add_heading(text, level=4)
        return heading

    # ========================================================================
    # BODY TEXT METHODS
    # ========================================================================

    def add_body_text(self, text, bold=False, italic=False, alignment=None):
        """
        Add a standard body paragraph.

        Args:
            text:      Paragraph text.
            bold:      Whether the text is bold.
            italic:    Whether the text is italic.
            alignment: Paragraph alignment (WD_ALIGN_PARAGRAPH).
        """
        para = self.doc.add_paragraph()
        if alignment:
            para.alignment = alignment
        run = para.add_run(text)
        run.font.name = Fonts.BODY
        run.font.size = FontSizes.BODY
        run.font.color.rgb = Colors.BODY_TEXT
        run.font.bold = bold
        run.font.italic = italic
        return para

    def add_body_bold(self, text):
        """Add a bold body paragraph."""
        return self.add_body_text(text, bold=True)

    def add_body_italic(self, text):
        """Add an italic body paragraph."""
        return self.add_body_text(text, italic=True)

    def add_small_text(self, text, color=None):
        """Add small-sized text, useful for captions and notes."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = Fonts.BODY
        run.font.size = FontSizes.SMALL
        run.font.color.rgb = color or Colors.LIGHT_TEXT
        return para

    def add_mixed_paragraph(self, parts):
        """
        Add a paragraph with mixed formatting (bold, italic, normal segments).

        Args:
            parts: List of tuples (text, bold, italic).
                   Example: [("Normal text ", False, False),
                             ("bold text", True, False),
                             (" and italic.", False, True)]
        """
        para = self.doc.add_paragraph()
        for text, bold, italic in parts:
            run = para.add_run(text)
            run.font.name = Fonts.BODY
            run.font.size = FontSizes.BODY
            run.font.color.rgb = Colors.BODY_TEXT
            run.font.bold = bold
            run.font.italic = italic
        return para

    # ========================================================================
    # LIST METHODS
    # ========================================================================

    def add_bullet_point(self, text, bold_prefix=""):
        """
        Add a single bullet point.

        Args:
            text:        Bullet text.
            bold_prefix: Optional bold text at the start (e.g., a label).
        """
        para = self.doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Spacing.BEFORE_LIST
        para.paragraph_format.space_after = Spacing.AFTER_LIST
        if bold_prefix:
            run_b = para.add_run(bold_prefix)
            run_b.font.name = Fonts.BODY
            run_b.font.size = FontSizes.BODY
            run_b.font.color.rgb = Colors.BODY_TEXT
            run_b.font.bold = True
        run = para.add_run(text)
        run.font.name = Fonts.BODY
        run.font.size = FontSizes.BODY
        run.font.color.rgb = Colors.BODY_TEXT
        return para

    def add_bullet_list(self, items):
        """
        Add multiple bullet points from a list of strings.

        Args:
            items: List of string items.
        """
        for item in items:
            self.add_bullet_point(item)

    def add_numbered_list(self, items):
        """
        Add a numbered list using List Number style.

        Args:
            items: List of string items.
        """
        for item in items:
            para = self.doc.add_paragraph(style='List Number')
            para.paragraph_format.space_before = Spacing.BEFORE_LIST
            para.paragraph_format.space_after = Spacing.AFTER_LIST
            run = para.add_run(item)
            run.font.name = Fonts.BODY
            run.font.size = FontSizes.BODY
            run.font.color.rgb = Colors.BODY_TEXT

    def add_keyed_bullet_list(self, items):
        """
        Add bullet points with bold keys.

        Args:
            items: List of tuples (key, description).
                   Example: [("Objective", "Build trust with investors.")]
        """
        for key, desc in items:
            self.add_bullet_point(f"  {desc}", bold_prefix=f"{key}:")

    # ========================================================================
    # TABLE METHODS
    # ========================================================================

    def add_professional_table(self, headers, rows, col_widths=None):
        """
        Add a professionally styled table with navy header and alternating rows.

        Args:
            headers:    List of header strings.
            rows:       List of lists (each inner list is a row of cell values).
            col_widths: Optional list of Inches values for column widths.
        """
        num_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # --- Set column widths if provided ---
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        # --- Header row ---
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            run.font.name = Fonts.BODY
            run.font.size = FontSizes.TABLE_HEADER
            run.font.color.rgb = Colors.WHITE_TEXT
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Navy background
            self._set_cell_shading(cell, Colors.TABLE_HEADER_BG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_padding(cell)

        # --- Data rows ---
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(cell_text))
                run.font.name = Fonts.BODY
                run.font.size = FontSizes.TABLE_BODY
                run.font.color.rgb = Colors.BODY_TEXT
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_cell_padding(cell)
                # Alternating row shading
                if row_idx % 2 == 1:
                    self._set_cell_shading(cell, Colors.TABLE_ALT_ROW_BG)

        # --- Table borders ---
        self._set_table_borders(table)

        # Spacer after table
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return table

    def add_info_table(self, data, key_width=2.2, val_width=4.0):
        """
        Add a two-column key-value information table.

        Args:
            data:      List of tuples (key, value).
            key_width: Width of the key column in inches.
            val_width: Width of the value column in inches.
        """
        table = self.doc.add_table(rows=len(data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for row_idx, (key, value) in enumerate(data):
            row = table.rows[row_idx]

            # Key cell
            key_cell = row.cells[0]
            key_cell.text = ""
            key_cell.width = Inches(key_width)
            kp = key_cell.paragraphs[0]
            kr = kp.add_run(key)
            kr.font.name = Fonts.BODY
            kr.font.size = FontSizes.TABLE_BODY
            kr.font.color.rgb = Colors.NAVY
            kr.font.bold = True
            self._set_cell_shading(key_cell, Colors.TABLE_ALT_ROW_BG)
            self._set_cell_padding(key_cell)

            # Value cell
            val_cell = row.cells[1]
            val_cell.text = ""
            val_cell.width = Inches(val_width)
            vp = val_cell.paragraphs[0]
            vr = vp.add_run(str(value))
            vr.font.name = Fonts.BODY
            vr.font.size = FontSizes.TABLE_BODY
            vr.font.color.rgb = Colors.BODY_TEXT
            self._set_cell_padding(val_cell)

        self._set_table_borders(table)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return table

    def add_checklist_table(self, headers, rows):
        """
        Add a checklist-style table with status column formatting.

        Args:
            headers: List of header strings (first should be "Item" or similar,
                     should include a "Status" column).
            rows:    List of lists.
        """
        return self.add_professional_table(headers, rows)

    def _set_cell_shading(self, cell, color_hex):
        """Apply background shading to a table cell."""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_padding(self, cell, top=60, bottom=60, left=80, right=80):
        """Set cell padding (in twips: 1 twip = 1/1440 inch)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for edge, val in [('top', top), ('bottom', bottom),
                          ('start', left), ('end', right)]:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)

    def _set_table_borders(self, table):
        """Apply consistent borders to a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), Colors.TABLE_BORDER_COLOR)
            borders.append(el)
        tblPr.append(borders)

    # ========================================================================
    # CALLOUT BOX METHODS
    # ========================================================================

    def add_callout_box(self, title, content, box_type="note"):
        """
        Add a styled callout box with left border accent and background shading.

        Args:
            title:    Callout box title.
            content:  Callout body text (string or list of strings).
            box_type: One of "note", "recommendation", "warning",
                      "future", "important".
        """
        # Resolve colors based on type
        config = {
            "note":           (Colors.NOTE_BG, Colors.NOTE_BORDER, Colors.NOTE_TITLE_COLOR, "\u2139\ufe0f  "),
            "recommendation": (Colors.REC_BG, Colors.REC_BORDER, Colors.REC_TITLE_COLOR, "\u2705  "),
            "warning":        (Colors.WARN_BG, Colors.WARN_BORDER, Colors.WARN_TITLE_COLOR, "\u26a0\ufe0f  "),
            "future":         (Colors.FUTURE_BG, Colors.FUTURE_BORDER, Colors.FUTURE_TITLE_COLOR, "\U0001f52e  "),
            "important":      (Colors.IMPORTANT_BG, Colors.IMPORTANT_BORDER, Colors.IMPORTANT_TITLE_COLOR, "\u2757  "),
        }
        bg_color, border_color, title_color, icon = config.get(
            box_type, config["note"]
        )

        # --- Title paragraph ---
        title_para = self.doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(10)
        title_para.paragraph_format.space_after = Pt(2)
        self._apply_callout_formatting(title_para, bg_color, border_color)
        run = title_para.add_run(f"{icon}{title.upper()}")
        run.font.name = Fonts.BODY
        run.font.size = FontSizes.CALLOUT_TITLE
        run.font.color.rgb = title_color
        run.font.bold = True

        # --- Content paragraph(s) ---
        if isinstance(content, list):
            for item in content:
                body_para = self.doc.add_paragraph()
                body_para.paragraph_format.space_before = Pt(1)
                body_para.paragraph_format.space_after = Pt(2)
                self._apply_callout_formatting(body_para, bg_color, border_color)
                run = body_para.add_run(f"\u2022  {item}")
                run.font.name = Fonts.BODY
                run.font.size = FontSizes.CALLOUT_BODY
                run.font.color.rgb = Colors.BODY_TEXT
        else:
            body_para = self.doc.add_paragraph()
            body_para.paragraph_format.space_before = Pt(1)
            body_para.paragraph_format.space_after = Pt(8)
            self._apply_callout_formatting(body_para, bg_color, border_color)
            run = body_para.add_run(content)
            run.font.name = Fonts.BODY
            run.font.size = FontSizes.CALLOUT_BODY
            run.font.color.rgb = Colors.BODY_TEXT

        # Add last-paragraph spacing
        if isinstance(content, list):
            body_para.paragraph_format.space_after = Pt(8)

    def add_note_box(self, title, content):
        """Shortcut for a Note callout box."""
        self.add_callout_box(title, content, box_type="note")

    def add_recommendation_box(self, title, content):
        """Shortcut for a Recommendation callout box."""
        self.add_callout_box(title, content, box_type="recommendation")

    def add_warning_box(self, title, content):
        """Shortcut for a Warning callout box."""
        self.add_callout_box(title, content, box_type="warning")

    def add_future_scope_box(self, title, content):
        """Shortcut for a Future Scope callout box."""
        self.add_callout_box(title, content, box_type="future")

    def add_important_box(self, title, content):
        """Shortcut for an Important callout box."""
        self.add_callout_box(title, content, box_type="important")

    def _apply_callout_formatting(self, paragraph, bg_color_hex, border_color_hex):
        """Apply left-border accent and background shading to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()

        # Left border
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), border_color_hex)
        pBdr.append(left)
        pPr.append(pBdr)

        # Background shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color_hex)
        pPr.append(shd)

        # Indentation for padding effect
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        ind.set(qn('w:right'), '360')
        pPr.append(ind)

    # ========================================================================
    # STRUCTURAL ELEMENTS
    # ========================================================================

    def add_page_break(self):
        """Insert an explicit page break."""
        self.doc.add_page_break()

    def add_section_separator(self):
        """Add a visual section separator (gold line)."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run("\u2500" * 50)
        run.font.color.rgb = Colors.GOLD
        run.font.size = Pt(10)

    def add_spacer(self, height_pt=12):
        """Add vertical whitespace."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(height_pt)
        return para

    # ========================================================================
    # PAGE SPECIFICATION HELPER
    # ========================================================================

    def add_page_spec(self, page_name, purpose, business_goal, user_goal,
                      sections, seo_info=None, future_scope=None,
                      recommendations=None, dependencies=None,
                      risks=None, assumptions=None):
        """
        Add a complete page specification block — reusable across documents.

        Args:
            page_name:       Name of the page.
            purpose:         Purpose statement.
            business_goal:   Business goal.
            user_goal:       User/visitor goal.
            sections:        List of tuples (section_name, description).
            seo_info:        Dict with keys like 'title', 'description', 'keywords', 'h1'.
            future_scope:    List of future scope items.
            recommendations: List of recommendation strings.
            dependencies:    List of dependency strings.
            risks:           List of risk strings.
            assumptions:     List of assumption strings.
        """
        self.add_heading_2(f"Page: {page_name}")

        # Info table
        info = [
            ("Page Name", page_name),
            ("Purpose", purpose),
            ("Business Goal", business_goal),
            ("User Goal", user_goal),
        ]
        self.add_info_table(info)

        # Sections
        if sections:
            self.add_heading_3("Section Breakdown")
            for sec_name, sec_desc in sections:
                self.add_heading_4(sec_name)
                self.add_body_text(sec_desc)

        # SEO
        if seo_info:
            self.add_heading_3("SEO Specification")
            seo_rows = [(k, v) for k, v in seo_info.items()]
            self.add_info_table(seo_rows)

        # Future Scope
        if future_scope:
            self.add_future_scope_box("Future Scope", future_scope)

        # Recommendations
        if recommendations:
            self.add_recommendation_box("Recommendations", recommendations)

        # Dependencies
        if dependencies:
            self.add_heading_3("Dependencies")
            self.add_bullet_list(dependencies)

        # Risks
        if risks:
            self.add_warning_box("Risks", risks)

        # Assumptions
        if assumptions:
            self.add_heading_3("Assumptions")
            self.add_bullet_list(assumptions)

    # ========================================================================
    # SAVE
    # ========================================================================

    def save(self, filename):
        """
        Save the document to the output directory.

        Args:
            filename: Output filename (e.g., "doc01_information_architecture.docx").

        Returns:
            Full path to the saved file.
        """
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        filepath = os.path.join(self.output_dir, filename)
        self.doc.save(filepath)
        print(f"  [OK] Saved: {filepath}")
        return filepath
